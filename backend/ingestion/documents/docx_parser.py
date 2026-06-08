import docx
import os

def extract_docx_text(file_path: str) -> str:
    """
    Extracts text from a DOCX file using python-docx, including paragraph text and tables.
    """
    if not os.path.exists(file_path):
        return ""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Extract tables as formatted pipe-separated lines
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n".join(full_text).strip()
    except Exception as e:
        print(f"python-docx failed for {file_path}: {e}")
        return ""
